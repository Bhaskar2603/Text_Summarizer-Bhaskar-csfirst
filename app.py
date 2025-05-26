from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, make_response
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
from heapq import nlargest
import nltk
import os
from docx import Document
import time
import PyPDF2
import pdfplumber
from PIL import Image
import io
from langdetect import detect, LangDetectException
from werkzeug.utils import secure_filename
import tempfile
import shutil
import io
from docx import Document as DocxDocument
from fpdf import FPDF
from datetime import datetime
import yake
import re
from nltk.tokenize.treebank import TreebankWordDetokenizer
from collections import defaultdict
from textblob import TextBlob
import nltk
nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')

def extract_keywords(text, num_keywords=5):
    """Extract keywords using YAKE! algorithm."""
    if not text or not text.strip():
        return []
        
    try:
        # Configure YAKE!
        language = "en"
        max_ngram_size = 3
        deduplication_threshold = 0.9
        
        # Initialize YAKE keyword extractor
        kw_extractor = yake.KeywordExtractor(
            lan=language,
            n=max_ngram_size,
            dedupLim=deduplication_threshold,
            top=num_keywords,
            features=None
        )
        
        # Extract keywords
        keywords = kw_extractor.extract_keywords(text)
        
        # Return only the keyword phrases (first element of each tuple)
        return [kw[0] for kw in keywords]
        
    except Exception as e:
        print(f"Error in YAKE keyword extraction: {str(e)}")
        # Fallback to simple word frequency if YAKE fails
        words = [word.lower() for word in word_tokenize(text) 
                if word.isalnum() and word.lower() not in stopwords.words('english')]
        freq_dist = FreqDist(words)
        return [word for word, _ in freq_dist.most_common(num_keywords)]

# Ensure all required NLTK data is available
def download_nltk_data():
    # Download punkt tokenizer for multiple languages
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        print("Downloading punkt tokenizer...")
        nltk.download('punkt')
    
    # Download stopwords for all supported languages
    supported_languages = [
        'english', 'spanish', 'french', 'german', 'italian', 
        'portuguese', 'russian', 'arabic', 'hindi', 'bengali',
        'chinese', 'japanese', 'korean'
    ]
    
    try:
        nltk.data.find('corpora/stopwords')
        print("Downloading stopwords...")
        nltk.download('stopwords')
    except Exception as e:
        print(f"Error downloading stopwords: {e}")
    
    # Ensure wordnet is available for better tokenization
    try:
        nltk.data.find('corpora/wordnet')
    except LookupError:
        print("Downloading wordnet...")
        nltk.download('wordnet')

# Download NLTK data
download_nltk_data()

app = Flask(__name__)
# Configure upload settings
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = os.path.join(tempfile.gettempdir(), 'text_summarizer_uploads')
app.config['ALLOWED_EXTENSIONS'] = {'txt', 'doc', 'docx'}

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/summarize', methods=['POST'])
def summarize():
    text = request.form['text']
    num_sentences = int(request.form['sentences'])
    
    if not text:
        return render_template('index.html', error="Please enter some text")
    
    summary = generate_summary(text, num_sentences)
    keywords = extract_keywords(text)
    return render_template('index.html', summary=summary, submitted_text=text, selected_sentences=num_sentences, keywords=keywords)

@app.route('/summarize_file', methods=['POST'])
def summarize_file():
    if 'file' not in request.files:
        return render_template('index.html', error="No file uploaded")
    
    file = request.files['file']
    if file.filename == '':
        return render_template('index.html', error="No file selected")
    
    # Only allow .txt, .doc, and .docx files
    if not allowed_file(file.filename):
        return render_template('index.html', 
                            error="Invalid file type. Only .txt, .doc, and .docx files are allowed")
    
    try:
        # Create a secure filename
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        # Ensure the upload directory exists
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # Save the file temporarily
        file.save(file_path)
        
        # Read file content based on file type
        text = ""
        try:
            if filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            elif filename.endswith(('.doc', '.docx')):
                doc = Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Clean up the temporary file
            if os.path.exists(file_path):
                os.remove(file_path)
                
            if not text.strip():
                return render_template('index.html', error="The file appears to be empty")
            
            # Get number of sentences from form or use default (3)
            num_sentences = int(request.form.get('sentences', 3))
            
            # Generate summary and extract keywords
            summary = generate_summary(text, num_sentences)
            keywords = extract_keywords(text)
            
            return render_template('index.html', 
                                 summary=summary,
                                 submitted_text=text,
                                 selected_sentences=num_sentences,
                                 keywords=keywords,
                                 active_tab='file-upload')
            
        except Exception as e:
            if os.path.exists(file_path):
                os.remove(file_path)
            return render_template('index.html', error=f"Error processing file: {str(e)}")
            
    except Exception as e:
        return render_template('index.html', error=f"Error uploading file: {str(e)}")

@app.route('/export_docx', methods=['POST'])
def export_docx():
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'No text provided'}), 400
        
        # Create a new Document
        doc = DocxDocument()
        
        # Add a title
        doc.add_heading('Text Summary', level=1)
        
        # Add the summary text directly
        doc.add_paragraph(data['text'])
        
        # Save the document to a BytesIO object
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Create a response with the file
        response = make_response(file_stream.getvalue())
        response.mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = 'attachment; filename=summary.docx'
        return response
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/summarize_pdf', methods=['POST'])
def summarize_pdf():
    if 'file' not in request.files:
        return render_template('index.html', error="No file uploaded")

    file = request.files['file']
    
    if file.filename == '':
        return render_template('index.html', error="No file selected")

    # Get file extension
    file_extension = os.path.splitext(file.filename)[1].lower()
    
    # Validate file type
    if file_extension not in ['.txt', '.docx', '.pdf']:
        return render_template('index.html', error="Unsupported file format. Only .txt, .docx, and .pdf files are allowed.")

    # Create uploads folder if it doesn't exist
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])

    try:
        # Create a unique filename to avoid conflicts
        filename = f"upload_{int(time.time())}{file_extension}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        # Save the file
        file.save(file_path)

        # Read file content
        text = ""
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif file_extension == '.docx':
            doc = Document(file_path)
            text = ' '.join([para.text for para in doc.paragraphs])
        elif file_extension == '.pdf':
            # Try PyPDF2 first
            try:
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
            except:
                # If PyPDF2 fails, try pdfplumber
                try:
                    with pdfplumber.open(file_path) as pdf:
                        text = ""
                        for page in pdf.pages:
                            text += page.extract_text()
                except:
                    return render_template('index.html', error="Unable to extract text from PDF. Please try a different file.")

        if not text.strip():
            return render_template('index.html', error="The file appears to be empty")

        # Get number of sentences from form
        num_sentences = int(request.form.get('file_sentences', 3))
        
        # Generate summary and extract keywords
        summary = generate_summary(text, num_sentences)
        keywords = extract_keywords(text)
        
        # Clean up uploaded file
        os.remove(file_path)
        
        return render_template('index.html', 
                             summary=summary, 
                             keywords=keywords,
                             active_tab='file-upload')
        
    except Exception as e:
        # Clean up uploaded file in case of error
        if os.path.exists(file_path):
            os.remove(file_path)
        return render_template('index.html', error=f"Error processing file: {str(e)}")

def detect_language(text):
    """Detect the language of the given text."""
    try:
        return detect(text)
    except LangDetectException:
        return 'en'  # Default to English if detection fails

def get_stopwords(lang):
    """Get stopwords for the specified language if available, else return empty set."""
    try:
        # Map language codes to NLTK stopwords names
        lang_map = {
            'en': 'english',
            'es': 'spanish',
            'fr': 'french',
            'de': 'german',
            'it': 'italian',
            'pt': 'portuguese',
            'ru': 'russian',
            'ar': 'arabic',
            'hi': 'hindi',
            'bn': 'bengali',
            'zh-cn': 'chinese',
            'ja': 'japanese',
            'ko': 'korean'
        }
        
        if lang in lang_map:
            try:
                return set(stopwords.words(lang_map[lang]))
            except (LookupError, OSError):
                # If stopwords for the language are not available, try to download them
                try:
                    nltk.download('stopwords')
                    return set(stopwords.words(lang_map[lang]))
                except:
                    return set()
        return set()
    except Exception as e:
        print(f"Error getting stopwords: {e}")
        return set()

def generate_summary(text, num_sentences):
    # Detect the language of the text
    lang = detect_language(text)
    print(f"Detected language: {lang}")
    
    # Get appropriate stopwords for the detected language
    stop_words = get_stopwords(lang)
    
    # Tokenize the text into sentences
    # Try to use language-specific tokenizer if available
    try:
        nltk.download('punkt')
        sentences = sent_tokenize(text, language='english' if lang == 'en' else lang)
    except:
        # Fallback to default tokenizer
        sentences = sent_tokenize(text)
    
    if not sentences:
        return text  # Return original if no sentences found
    
    # Tokenize the text into words and remove stopwords
    words = [word.lower() for word in word_tokenize(text) if word.isalnum()]
    words = [word for word in words if word not in stop_words]
    
    # Calculate word frequencies
    freq_dist = FreqDist(words)
    
    # Calculate sentence scores based on word frequencies
    sentence_scores = {}
    for sentence in sentences:
        for word in word_tokenize(sentence.lower()):
            if word in freq_dist:
                if sentence not in sentence_scores:
                    sentence_scores[sentence] = freq_dist[word]
                else:
                    sentence_scores[sentence] += freq_dist[word]
    
    # Get the top N sentences with highest scores
    # If we have fewer sentences than requested, return what we have
    num_sentences = min(num_sentences, len(sentences))
    summary_sentences = nlargest(num_sentences, sentence_scores, key=sentence_scores.get)
    
    # Join the summary sentences
    summary = ' '.join(summary_sentences)
    return summary

@app.route('/process_voice_input', methods=['POST'])
def process_voice_input():
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Clean and format the text
        text = text.capitalize()
        
        # Use TextBlob to improve the text
        blob = TextBlob(text)
        
        # Get the corrected text
        corrected_text = str(blob.correct())
        
        # Add proper capitalization and punctuation
        sentences = nltk.sent_tokenize(corrected_text)
        formatted_sentences = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # Capitalize first letter
            sentence = sentence[0].upper() + sentence[1:]
            
            # Ensure it ends with punctuation
            if not sentence[-1] in '.!?':
                sentence += '.'
                
            formatted_sentences.append(sentence)
        
        processed_text = ' '.join(formatted_sentences)
        
        return jsonify({
            'original_text': text,
            'processed_text': processed_text
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    host = os.getenv('FLASK_RUN_HOST', '0.0.0.0')
    port = int(os.getenv('FLASK_RUN_PORT', 5000))
    app.run(host=host, port=port, debug=True)
